"""
Complete DevOps Setup Guide Generator
Run: pip install python-docx
Then: python docs/generate_complete_guide.py
Output: docs/Complete_DevOps_Setup_Guide_FA22BSE100.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "Complete_DevOps_Setup_Guide_FA22BSE100.docx")

# ── Colours ──────────────────────────────────────────────────────────────────
ORANGE  = RGBColor(0xE8, 0x50, 0x00)
BLUE    = RGBColor(0x1F, 0x49, 0x7D)
GREEN   = RGBColor(0x10, 0x7C, 0x10)
GRAY    = RGBColor(0x36, 0x36, 0x36)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
YELLOW_BG = 'FFF9C4'
GREEN_BG  = 'E8F5E9'
RED_BG    = 'FFEBEE'
BLUE_BG   = 'E3F2FD'
CODE_BG   = 'F4F4F4'

def shd(cell, fill):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    s  = OxmlElement('w:shd')
    s.set(qn('w:val'),   'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'),  fill)
    pr.append(s)

def pgshd(p, fill):
    pr = p._p.get_or_add_pPr()
    s  = OxmlElement('w:shd')
    s.set(qn('w:val'),   'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'),  fill)
    pr.append(s)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(doc, text):
    p = doc.add_heading('', level=1)
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = ORANGE
    run.font.size = Pt(18)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(doc, text):
    p = doc.add_heading('', level=2)
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = BLUE
    run.font.size = Pt(14)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(doc, text):
    p = doc.add_heading('', level=3)
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = GREEN
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(5)
    for r in p.runs:
        r.font.size = Pt(11)
        r.font.color.rgb = GRAY
    return p

def code(doc, text, label=None):
    if label:
        lp = doc.add_paragraph()
        lr = lp.add_run(f'  {label}')
        lr.font.size  = Pt(8)
        lr.font.bold  = True
        lr.font.color.rgb = RGBColor(0x66,0x66,0x66)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    pgshd(p, CODE_BG)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x8C)
    return p

def note(doc, text, bg=YELLOW_BG, prefix='NOTE'):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]
    shd(c, bg)
    p = c.paragraphs[0]
    r1 = p.add_run(f'{prefix}: ')
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x5D,0x40,0x00) if bg == YELLOW_BG else RGBColor(0x1B,0x5E,0x20)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()

def warn(doc, text):
    note(doc, text, RED_BG, 'WARNING')

def tip(doc, text):
    note(doc, text, GREEN_BG, 'TIP')

def info(doc, text):
    note(doc, text, BLUE_BG, 'INFO')

def steps(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY

def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY

def table2(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Shading Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def pagebreak(doc):
    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════════
def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin    = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin   = Cm(2.5)
        s.right_margin  = Cm(2.5)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Complete DevOps Pipeline')
    r.font.size = Pt(32); r.font.bold = True; r.font.color.rgb = ORANGE

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run('Setup Guide — From Zero to Full Pipeline')
    r2.font.size = Pt(16); r2.font.color.rgb = BLUE

    doc.add_paragraph()
    stack = doc.add_paragraph()
    stack.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = stack.add_run(
        'AWS EC2  •  Terraform  •  Ansible  •  Docker  •  Jenkins\n'
        'Kubernetes (Minikube)  •  SonarQube  •  Prometheus  •  Grafana'
    )
    rs.font.size = Pt(12); rs.font.color.rgb = GRAY

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(
        'Muhammad Wahaj Naveed\nFA22-BSE-100\nMay 2026'
    )
    rm.font.size = Pt(12); rm.font.color.rgb = GRAY

    doc.add_paragraph()
    arch = doc.add_paragraph()
    arch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ra = arch.add_run(
        'GitHub → Jenkins CI/CD → SonarQube Quality Gate\n'
        '→ Docker Build → Minikube Kubernetes Pod\n'
        '← Prometheus + Grafana Monitoring'
    )
    ra.font.size = Pt(11); ra.font.color.rgb = BLUE; ra.font.italic = True

    pagebreak(doc)

    # ── TABLE OF CONTENTS (manual) ────────────────────────────────────────────
    h1(doc, 'Table of Contents')
    toc = [
        ('1',  'Prerequisites & Tools Required'),
        ('2',  'Architecture Overview'),
        ('3',  'AWS EC2 Instance Setup'),
        ('4',  'Infrastructure as Code — Terraform'),
        ('5',  'System Setup on EC2'),
        ('6',  'Docker Installation'),
        ('7',  'Jenkins Installation & Configuration'),
        ('8',  'Ansible — Configuration Management'),
        ('9',  'Kubernetes with Minikube'),
        ('10', 'Prometheus & Node Exporter — Metrics Collection'),
        ('11', 'Grafana — Monitoring Dashboards'),
        ('12', 'SonarQube — Code Quality & Security'),
        ('13', 'Jenkins Pipeline — Full CI/CD Setup'),
        ('14', 'Complete End-to-End Flow'),
        ('15', 'Service URLs Reference'),
        ('16', 'Troubleshooting'),
    ]
    for num, title in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f'  Part {num}:  {title}')
        r.font.size = Pt(11); r.font.color.rgb = BLUE

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 1 — PREREQUISITES
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 1: Prerequisites & Tools Required')
    body(doc, 'Before starting, make sure you have the following accounts and tools on your local Windows machine:')

    h2(doc, '1.1 Accounts Needed')
    table2(doc,
        ['Account', 'Purpose', 'URL'],
        [
            ('AWS Account',    'Host the EC2 instance',        'aws.amazon.com'),
            ('GitHub Account', 'Source code repository',       'github.com'),
            ('SonarCloud (optional)', 'Alternative to self-hosted SonarQube', 'sonarcloud.io'),
        ]
    )

    h2(doc, '1.2 Local Tools to Install')
    table2(doc,
        ['Tool', 'Version', 'Purpose'],
        [
            ('Git',            'Latest',   'Push code to GitHub'),
            ('VS Code',        'Latest',   'Code editor + SonarLint plugin'),
            ('Terraform',      '>= 1.0',   'Provision AWS infrastructure'),
            ('AWS CLI',        'v2',       'Configure AWS credentials'),
            ('Python 3',       '>= 3.8',   'Run doc generation script'),
            ('python-docx',    'Latest',   'pip install python-docx'),
        ]
    )

    h2(doc, '1.3 AWS Key Pair')
    body(doc, 'You need an EC2 key pair (.pem file) to SSH into your instance.')
    steps(doc, [
        'AWS Console → EC2 → Key Pairs → Create key pair',
        'Name: jenkey | Type: RSA | Format: .pem',
        'Download and save to C:\\Users\\<you>\\Downloads\\jenkey.pem',
        'On Windows PowerShell: icacls jenkey.pem /inheritance:r /grant:r "%USERNAME%:R"',
    ])

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 2 — ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 2: Architecture Overview')
    body(doc, 'Everything runs on a single AWS EC2 instance (t2.medium, Ubuntu 24.04). The full pipeline flow is:')

    code(doc,
        'Developer (VS Code + SonarLint)\n'
        '        │\n'
        '        │  git push\n'
        '        ▼\n'
        'GitHub Repository (main branch)\n'
        '        │\n'
        '        │  webhook / manual trigger\n'
        '        ▼\n'
        '┌─────────────────────────────────────────────────────┐\n'
        '│              AWS EC2 (t2.medium, Ubuntu 24.04)      │\n'
        '│                                                     │\n'
        '│  Jenkins :8080                                      │\n'
        '│    ├── Stage 1: Clone Repo from GitHub              │\n'
        '│    ├── Stage 2: SonarQube Analysis                  │\n'
        '│    │       └── SonarQube :9000 (Docker container)   │\n'
        '│    ├── Stage 3: Quality Gate Check (BLOCK if FAIL)  │\n'
        '│    ├── Stage 4: Docker Build → wahaj-app:latest     │\n'
        '│    ├── Stage 5: Deploy Docker Container :3001       │\n'
        '│    ├── Stage 6: Load Image → Minikube               │\n'
        '│    ├── Stage 7: kubectl apply → K8s Pod :30080      │\n'
        '│    └── Stage 8: Verify Deployment                   │\n'
        '│                                                     │\n'
        '│  Prometheus :9090 ◄── Node Exporter :9100          │\n'
        '│  Grafana    :3030 ◄── Prometheus                    │\n'
        '└─────────────────────────────────────────────────────┘'
    )

    h2(doc, '2.1 Port Reference')
    table2(doc,
        ['Service', 'Port', 'Protocol', 'Description'],
        [
            ('SSH',            '22',        'TCP', 'Secure shell access'),
            ('Jenkins',        '8080',      'TCP', 'CI/CD web interface'),
            ('Next.js (Docker)','3001',     'TCP', 'App via Docker container'),
            ('Next.js (K8s)',   '30080',    'TCP', 'App via Kubernetes NodePort'),
            ('SonarQube',      '9000',      'TCP', 'Code quality dashboard'),
            ('Prometheus',     '9090',      'TCP', 'Metrics database UI'),
            ('Node Exporter',  '9100',      'TCP', 'System metrics endpoint'),
            ('Grafana',        '3030',      'TCP', 'Monitoring dashboards'),
        ]
    )

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 3 — AWS EC2
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 3: AWS EC2 Instance Setup')

    h2(doc, '3.1 Launch EC2 Instance')
    steps(doc, [
        'Log in to AWS Console → go to EC2 → click Launch Instance',
        'Name: wahaj-devops-server',
        'AMI: Ubuntu Server 24.04 LTS (Free tier eligible)',
        'Instance type: t2.medium (2 vCPU, 4 GB RAM) — minimum required',
        'Key pair: select jenkey (created in Part 1)',
        'Storage: 34 GB gp3 (increase from default 8 GB)',
        'Click Launch Instance',
    ])
    note(doc, 't2.micro does NOT have enough RAM for SonarQube + Minikube. Use t2.medium or larger.')

    h2(doc, '3.2 Configure Security Group (Open Ports)')
    body(doc, 'AWS Console → EC2 → Security Groups → select your group → Edit Inbound Rules → Add each row:')
    table2(doc,
        ['Port', 'Protocol', 'Source', 'Purpose'],
        [
            ('22',        'TCP', '0.0.0.0/0', 'SSH'),
            ('80',        'TCP', '0.0.0.0/0', 'HTTP'),
            ('3000',      'TCP', '0.0.0.0/0', 'Next.js app'),
            ('3001',      'TCP', '0.0.0.0/0', 'Next.js Docker alt port'),
            ('3030',      'TCP', '0.0.0.0/0', 'Grafana'),
            ('8080',      'TCP', '0.0.0.0/0', 'Jenkins'),
            ('9000',      'TCP', '0.0.0.0/0', 'SonarQube'),
            ('9090',      'TCP', '0.0.0.0/0', 'Prometheus'),
            ('9100',      'TCP', '0.0.0.0/0', 'Node Exporter'),
            ('30000-32767','TCP','0.0.0.0/0', 'Kubernetes NodePorts'),
        ]
    )

    h2(doc, '3.3 Connect to EC2 via SSH')
    body(doc, 'From PowerShell on your Windows machine:')
    code(doc,
        '# Replace <EC2-PUBLIC-IP> with your actual IP from AWS Console\n'
        'ssh -i C:\\Users\\<you>\\Downloads\\jenkey.pem ubuntu@<EC2-PUBLIC-IP>\n\n'
        '# Example:\n'
        'ssh -i C:\\Users\\wahaj\\Downloads\\jenkey.pem ubuntu@54.234.94.152',
        'PowerShell on Windows'
    )
    note(doc, 'EC2 public IP changes every time you stop/start the instance. Always get the current IP from AWS Console → EC2 → Instances → Public IPv4 address.')

    h2(doc, '3.4 Expand Filesystem After Storage Increase')
    body(doc, 'If you increase EBS volume size in AWS Console, run these on EC2 to claim the new space:')
    code(doc,
        'sudo growpart /dev/nvme0n1 1\n'
        'sudo resize2fs /dev/root\n'
        'df -h   # Verify new size shows up',
        'EC2 SSH'
    )

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 4 — TERRAFORM
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 4: Infrastructure as Code — Terraform')
    body(doc, 'Terraform automates EC2 and security group creation so the entire infrastructure is reproducible.')

    h2(doc, '4.1 Install Terraform (Windows)')
    steps(doc, [
        'Download from: https://developer.hashicorp.com/terraform/install',
        'Extract terraform.exe to C:\\Windows\\System32\\ (or add to PATH)',
        'Verify: open PowerShell → terraform --version',
    ])

    h2(doc, '4.2 Configure AWS CLI Credentials')
    code(doc,
        'aws configure\n'
        '# Enter when prompted:\n'
        '# AWS Access Key ID:     <your-access-key>\n'
        '# AWS Secret Access Key: <your-secret-key>\n'
        '# Default region:        us-east-1\n'
        '# Default output format: json',
        'PowerShell on Windows'
    )
    note(doc, 'Get your Access Key from AWS Console → IAM → Users → your user → Security credentials → Create access key.')

    h2(doc, '4.3 Terraform Files Structure')
    code(doc,
        'terraform/\n'
        '├── main.tf        # EC2 instance + security group resources\n'
        '├── variables.tf   # Configurable inputs\n'
        '└── outputs.tf     # Printed URLs after apply'
    )

    h2(doc, '4.4 main.tf — Key Configuration')
    code(doc,
        '# terraform/main.tf\n'
        'provider "aws" {\n'
        '  region = "us-east-1"\n'
        '}\n\n'
        'resource "aws_instance" "wahaj_ec2" {\n'
        '  ami           = "ami-0c7217cdde317cfec"  # Ubuntu 24.04\n'
        '  instance_type = "t2.medium"\n'
        '  key_name      = "jenkey"\n'
        '  root_block_device {\n'
        '    volume_size = 34\n'
        '    volume_type = "gp3"\n'
        '  }\n'
        '}',
        'terraform/main.tf'
    )

    h2(doc, '4.5 Run Terraform')
    code(doc,
        'cd terraform\n'
        'terraform init       # Download AWS provider plugin\n'
        'terraform plan       # Preview: shows what will be created\n'
        'terraform apply      # Create EC2 instance and security group\n'
        '                     # Type "yes" when prompted\n'
        'terraform output     # Print all service URLs\n\n'
        '# To destroy everything when done:\n'
        'terraform destroy',
        'PowerShell on Windows'
    )

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 5 — SYSTEM SETUP
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 5: System Setup on EC2')
    body(doc, 'Run these first on every fresh EC2 instance before installing anything else.')

    h2(doc, '5.1 Update System & Install Base Packages')
    code(doc,
        'sudo apt-get update && sudo apt-get upgrade -y\n\n'
        'sudo apt-get install -y \\\n'
        '  curl wget git unzip \\\n'
        '  apt-transport-https ca-certificates \\\n'
        '  software-properties-common gnupg lsb-release \\\n'
        '  python3 python3-pip net-tools',
        'EC2 SSH'
    )

    h2(doc, '5.2 Set Kernel Parameters (Required for SonarQube/Elasticsearch)')
    code(doc,
        'sudo sysctl -w vm.max_map_count=524288\n'
        'sudo sysctl -w fs.file-max=131072\n'
        'echo "vm.max_map_count=524288" | sudo tee -a /etc/sysctl.conf\n'
        'echo "fs.file-max=131072"      | sudo tee -a /etc/sysctl.conf',
        'EC2 SSH'
    )

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 6 — DOCKER
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 6: Docker Installation')

    h2(doc, '6.1 Install Docker CE')
    code(doc,
        '# Add Docker GPG key\n'
        'curl -fsSL https://download.docker.com/linux/ubuntu/gpg | \\\n'
        '  sudo gpg --dearmor -o /usr/share/keyrings/docker-archive-keyring.gpg\n\n'
        '# Add Docker repository\n'
        'echo "deb [arch=amd64 signed-by=/usr/share/keyrings/docker-archive-keyring.gpg] \\\n'
        '  https://download.docker.com/linux/ubuntu $(lsb_release -cs) stable" | \\\n'
        '  sudo tee /etc/apt/sources.list.d/docker.list\n\n'
        '# Install Docker\n'
        'sudo apt-get update\n'
        'sudo apt-get install -y docker-ce docker-ce-cli containerd.io\n\n'
        '# Start and enable Docker\n'
        'sudo systemctl start docker\n'
        'sudo systemctl enable docker\n\n'
        '# Verify\n'
        'docker --version',
        'EC2 SSH'
    )

    h2(doc, '6.2 Add Users to Docker Group')
    code(doc,
        '# Add ubuntu user (for manual Docker commands)\n'
        'sudo usermod -aG docker ubuntu\n\n'
        '# Add jenkins user (for Jenkins pipeline)\n'
        'sudo usermod -aG docker jenkins\n\n'
        '# Apply group change in current session\n'
        'newgrp docker\n\n'
        '# Restart Jenkins to apply the group change\n'
        'sudo systemctl restart jenkins',
        'EC2 SSH'
    )
    note(doc, 'You must add BOTH ubuntu and jenkins users to the docker group. Without this, Docker commands will get "permission denied" errors.')

    h2(doc, '6.3 Test Docker')
    code(doc,
        'docker run hello-world\n'
        '# Should print "Hello from Docker!"',
        'EC2 SSH'
    )

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 7 — JENKINS
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 7: Jenkins Installation & Configuration')

    h2(doc, '7.1 Install Java 17 (Required by Jenkins)')
    code(doc,
        'sudo apt-get update\n'
        'sudo apt-get install -y openjdk-17-jdk\n'
        'java -version   # Should show: openjdk 17.x.x',
        'EC2 SSH'
    )

    h2(doc, '7.2 Install Jenkins')
    code(doc,
        '# Add Jenkins GPG key and repository\n'
        'curl -fsSL https://pkg.jenkins.io/debian-stable/jenkins.io-2023.key | \\\n'
        '  sudo tee /usr/share/keyrings/jenkins-keyring.asc > /dev/null\n\n'
        'echo "deb [signed-by=/usr/share/keyrings/jenkins-keyring.asc] \\\n'
        '  https://pkg.jenkins.io/debian-stable binary/" | \\\n'
        '  sudo tee /etc/apt/sources.list.d/jenkins.list\n\n'
        '# Install Jenkins\n'
        'sudo apt-get update\n'
        'sudo apt-get install -y jenkins\n\n'
        '# Start and enable Jenkins\n'
        'sudo systemctl start jenkins\n'
        'sudo systemctl enable jenkins\n'
        'sudo systemctl status jenkins',
        'EC2 SSH'
    )

    h2(doc, '7.3 Initial Jenkins Setup')
    steps(doc, [
        'Open browser: http://<EC2-IP>:8080',
        'Get initial admin password:  sudo cat /var/lib/jenkins/secrets/initialAdminPassword',
        'Paste the password into Jenkins and click Continue',
        'Click "Install suggested plugins" and wait',
        'Create admin user: Username=wahaj, set a password, save',
        'Click "Save and Finish" → "Start using Jenkins"',
    ])

    h2(doc, '7.4 Install Required Jenkins Plugins')
    steps(doc, [
        'Manage Jenkins → Plugins → Available plugins',
        'Search and install: SonarQube Scanner',
        'Search and install: Sonar Quality Gates',
        'Search and install: Docker Pipeline (optional but useful)',
        'Click Install and restart Jenkins when prompted',
    ])

    h2(doc, '7.5 Configure Global Tools in Jenkins')
    body(doc, 'Manage Jenkins → Tools:')
    steps(doc, [
        'JDK → Add JDK → Name: Java17 → Install automatically → Version: 17',
        'SonarQube Scanner → Add → Name: SonarQube → Install automatically',
        'Click Save',
    ])

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 8 — ANSIBLE
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 8: Ansible — Configuration Management')
    body(doc, 'Ansible automates configuration of all services. Run Ansible from your LOCAL Windows machine (not EC2). It connects to EC2 via SSH and configures everything automatically.')

    h2(doc, '8.1 Install Ansible on Windows (via WSL) or Linux')
    code(doc,
        '# On WSL (Windows Subsystem for Linux) or Ubuntu:\n'
        'sudo apt-get update\n'
        'sudo apt-get install -y ansible\n'
        'ansible --version',
        'WSL / Linux'
    )
    note(doc, 'Ansible cannot run natively on Windows. Use WSL (Windows Subsystem for Linux) or run Ansible directly from the EC2 instance against localhost.')

    h2(doc, '8.2 Inventory File')
    body(doc, 'Create ansible/inventory.ini — tells Ansible which server to connect to:')
    code(doc,
        '[ec2]\n'
        'wahaj-server ansible_host=<EC2-PUBLIC-IP> \\\n'
        '  ansible_user=ubuntu \\\n'
        '  ansible_ssh_private_key_file=~/.ssh/jenkey.pem \\\n'
        '  ansible_ssh_common_args=\'-o StrictHostKeyChecking=no\'',
        'ansible/inventory.ini'
    )
    warn(doc, 'Update ansible_host= every time your EC2 restarts and gets a new public IP.')

    h2(doc, '8.3 Test Connectivity')
    code(doc,
        'ansible -i ansible/inventory.ini ec2 -m ping\n'
        '# Expected: wahaj-server | SUCCESS => {"ping": "pong"}',
        'Local machine / WSL'
    )

    h2(doc, '8.4 Run Base Stack Playbook')
    body(doc, 'Installs Docker, kubectl, Minikube, Prometheus, Node Exporter, Grafana:')
    code(doc,
        'ansible-playbook -i ansible/inventory.ini ansible/playbook.yml\n\n'
        '# Run with verbose output if debugging:\n'
        'ansible-playbook -i ansible/inventory.ini ansible/playbook.yml -v',
        'Local machine / WSL'
    )

    h2(doc, '8.5 Run SonarQube Playbook')
    body(doc, 'Installs Java 17, PostgreSQL, SonarQube 10.4, SonarScanner CLI:')
    code(doc,
        'ansible-playbook -i ansible/inventory.ini ansible/sonarqube-playbook.yml',
        'Local machine / WSL'
    )

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 9 — KUBERNETES
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 9: Kubernetes with Minikube')

    h2(doc, '9.1 Install kubectl')
    code(doc,
        'curl -LO "https://dl.k8s.io/release/v1.28.0/bin/linux/amd64/kubectl"\n'
        'sudo install -o root -g root -m 0755 kubectl /usr/local/bin/kubectl\n'
        'kubectl version --client',
        'EC2 SSH'
    )

    h2(doc, '9.2 Install Minikube')
    code(doc,
        'curl -LO https://storage.googleapis.com/minikube/releases/latest/minikube-linux-amd64\n'
        'sudo install minikube-linux-amd64 /usr/local/bin/minikube\n'
        'minikube version',
        'EC2 SSH'
    )

    h2(doc, '9.3 Start Minikube')
    body(doc, 'IMPORTANT: Minikube must be started for the SAME USER that runs Jenkins (jenkins user):')
    code(doc,
        '# Start Minikube as the jenkins user\n'
        'sudo su - jenkins -s /bin/bash -c \\\n'
        '  "minikube start --driver=docker --memory=2048 --cpus=2"\n\n'
        '# Verify\n'
        'sudo su - jenkins -s /bin/bash -c "minikube status"\n'
        'sudo su - jenkins -s /bin/bash -c "kubectl get nodes"',
        'EC2 SSH'
    )
    note(doc, 'If you restart EC2, Minikube stops. Run the start command again before triggering Jenkins.')

    h2(doc, '9.4 Kubernetes Manifests')
    body(doc, 'These files live in the kubernetes/ directory of the repository:')
    code(doc,
        '# kubernetes/namespace.yaml\n'
        'apiVersion: v1\n'
        'kind: Namespace\n'
        'metadata:\n'
        '  name: wahaj-app\n\n'
        '# kubernetes/deployment.yaml (key settings)\n'
        'spec:\n'
        '  replicas: 1\n'
        '  containers:\n'
        '    - name: wahaj-app\n'
        '      image: wahaj-app:latest\n'
        '      imagePullPolicy: Never   # Use local Minikube image\n'
        '      ports:\n'
        '        - containerPort: 3000\n\n'
        '# kubernetes/service.yaml\n'
        'spec:\n'
        '  type: NodePort\n'
        '  ports:\n'
        '    - port: 3000\n'
        '      nodePort: 30080         # External access port',
        'kubernetes/'
    )

    h2(doc, '9.5 Manual Kubernetes Deploy Commands')
    code(doc,
        '# Apply manifests\n'
        'kubectl apply -f kubernetes/namespace.yaml\n'
        'kubectl apply -f kubernetes/deployment.yaml\n'
        'kubectl apply -f kubernetes/service.yaml\n\n'
        '# Check status\n'
        'kubectl get pods -n wahaj-app\n'
        'kubectl get svc  -n wahaj-app\n\n'
        '# View logs\n'
        'kubectl logs -f deployment/wahaj-app -n wahaj-app\n\n'
        '# Restart deployment (after new image load)\n'
        'kubectl rollout restart deployment/wahaj-app -n wahaj-app\n'
        'kubectl rollout status  deployment/wahaj-app -n wahaj-app',
        'EC2 SSH'
    )

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 10 — PROMETHEUS & NODE EXPORTER
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 10: Prometheus & Node Exporter')

    h2(doc, '10.1 Create System Users')
    code(doc,
        'sudo useradd --no-create-home --shell /bin/false prometheus\n'
        'sudo useradd --no-create-home --shell /bin/false node_exporter',
        'EC2 SSH'
    )

    h2(doc, '10.2 Install Prometheus')
    code(doc,
        'cd /tmp\n'
        'wget https://github.com/prometheus/prometheus/releases/download/\\\n'
        'v2.47.0/prometheus-2.47.0.linux-amd64.tar.gz\n'
        'tar -xvf prometheus-2.47.0.linux-amd64.tar.gz\n\n'
        '# Create directories\n'
        'sudo mkdir /etc/prometheus /var/lib/prometheus\n\n'
        '# Copy binaries\n'
        'sudo cp /tmp/prometheus-2.47.0.linux-amd64/prometheus  /usr/local/bin/\n'
        'sudo cp /tmp/prometheus-2.47.0.linux-amd64/promtool    /usr/local/bin/\n'
        'sudo cp -r /tmp/prometheus-2.47.0.linux-amd64/consoles          /etc/prometheus\n'
        'sudo cp -r /tmp/prometheus-2.47.0.linux-amd64/console_libraries /etc/prometheus\n\n'
        '# Set ownership\n'
        'sudo chown -R prometheus:prometheus /etc/prometheus /var/lib/prometheus\n'
        'sudo chown prometheus:prometheus /usr/local/bin/prometheus /usr/local/bin/promtool',
        'EC2 SSH'
    )

    h2(doc, '10.3 Prometheus Configuration File')
    code(doc,
        'sudo nano /etc/prometheus/prometheus.yml\n\n'
        '# Paste this content:\n'
        'global:\n'
        '  scrape_interval: 15s\n'
        '  evaluation_interval: 15s\n\n'
        'scrape_configs:\n'
        '  - job_name: prometheus\n'
        '    static_configs:\n'
        '      - targets: [\'localhost:9090\']\n\n'
        '  - job_name: node_exporter\n'
        '    static_configs:\n'
        '      - targets: [\'localhost:9100\']\n\n'
        '  - job_name: jenkins\n'
        '    metrics_path: /prometheus\n'
        '    static_configs:\n'
        '      - targets: [\'localhost:8080\']\n\n'
        'sudo chown prometheus:prometheus /etc/prometheus/prometheus.yml',
        'EC2 SSH'
    )

    h2(doc, '10.4 Prometheus Systemd Service')
    code(doc,
        'sudo nano /etc/systemd/system/prometheus.service\n\n'
        '# Paste:\n'
        '[Unit]\n'
        'Description=Prometheus Monitoring\n'
        'After=network-online.target\n\n'
        '[Service]\n'
        'User=prometheus\n'
        'Group=prometheus\n'
        'Type=simple\n'
        'ExecStart=/usr/local/bin/prometheus \\\n'
        '  --config.file=/etc/prometheus/prometheus.yml \\\n'
        '  --storage.tsdb.path=/var/lib/prometheus/ \\\n'
        '  --web.listen-address=0.0.0.0:9090\n\n'
        '[Install]\n'
        'WantedBy=multi-user.target\n\n'
        '# Start Prometheus\n'
        'sudo systemctl daemon-reload\n'
        'sudo systemctl start prometheus\n'
        'sudo systemctl enable prometheus\n'
        'sudo systemctl status prometheus',
        'EC2 SSH'
    )

    h2(doc, '10.5 Install Node Exporter')
    code(doc,
        'cd /tmp\n'
        'wget https://github.com/prometheus/node_exporter/releases/download/\\\n'
        'v1.6.1/node_exporter-1.6.1.linux-amd64.tar.gz\n'
        'tar -xvf node_exporter-1.6.1.linux-amd64.tar.gz\n'
        'sudo cp /tmp/node_exporter-1.6.1.linux-amd64/node_exporter /usr/local/bin/\n'
        'sudo chown node_exporter:node_exporter /usr/local/bin/node_exporter\n\n'
        'sudo nano /etc/systemd/system/node_exporter.service\n\n'
        '# Paste:\n'
        '[Unit]\n'
        'Description=Node Exporter\n'
        'After=network-online.target\n\n'
        '[Service]\n'
        'User=node_exporter\n'
        'ExecStart=/usr/local/bin/node_exporter\n\n'
        '[Install]\n'
        'WantedBy=multi-user.target\n\n'
        'sudo systemctl daemon-reload\n'
        'sudo systemctl start node_exporter\n'
        'sudo systemctl enable node_exporter\n'
        'sudo systemctl status node_exporter',
        'EC2 SSH'
    )
    tip(doc, 'Verify Node Exporter: open http://<EC2-IP>:9100/metrics in browser — you should see a long list of system metrics.')

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 11 — GRAFANA
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 11: Grafana — Monitoring Dashboards')

    h2(doc, '11.1 Install Grafana')
    code(doc,
        'sudo apt-get install -y apt-transport-https software-properties-common wget\n\n'
        '# Add Grafana GPG key\n'
        'sudo mkdir -p /etc/apt/keyrings\n'
        'wget -q -O - https://apt.grafana.com/gpg.key | \\\n'
        '  gpg --dearmor | sudo tee /etc/apt/keyrings/grafana.gpg > /dev/null\n\n'
        '# Add repository\n'
        'echo "deb [signed-by=/etc/apt/keyrings/grafana.gpg] \\\n'
        '  https://apt.grafana.com stable main" | \\\n'
        '  sudo tee /etc/apt/sources.list.d/grafana.list\n\n'
        'sudo apt-get update\n'
        'sudo apt-get install -y grafana',
        'EC2 SSH'
    )

    h2(doc, '11.2 Change Grafana Port to 3030')
    body(doc, 'Port 3000 is used by the Next.js app, so change Grafana to 3030:')
    code(doc,
        'sudo sed -i \'/^\\[server\\]/a http_port = 3030\' /etc/grafana/grafana.ini\n\n'
        '# Or edit manually:\n'
        'sudo nano /etc/grafana/grafana.ini\n'
        '# Find [server] section and set: http_port = 3030\n\n'
        'sudo systemctl start grafana-server\n'
        'sudo systemctl enable grafana-server\n'
        'sudo systemctl status grafana-server',
        'EC2 SSH'
    )

    h2(doc, '11.3 First Login & Add Prometheus Data Source')
    steps(doc, [
        'Open http://<EC2-IP>:3030 — login: admin / admin',
        'Change password when prompted',
        'Left sidebar → Connections → Data Sources → Add data source',
        'Select Prometheus',
        'URL: http://localhost:9090',
        'Click Save & Test → should show green "Data source is working"',
    ])

    h2(doc, '11.4 Import Pre-Built Dashboards')
    steps(doc, [
        'Left sidebar → Dashboards → Import',
        'Enter ID 1860 → Load → select Prometheus → Import (Node Exporter Full — CPU/RAM/Disk)',
        'Repeat: Import ID 9964 → (Jenkins Performance and Health Overview)',
    ])

    h2(doc, '11.5 Enable Jenkins Prometheus Plugin')
    steps(doc, [
        'Jenkins → Manage Jenkins → Plugins → Available → search Prometheus metrics → Install',
        'After install, verify: http://<EC2-IP>:8080/prometheus (shows Jenkins metrics)',
        'Prometheus will now scrape Jenkins metrics automatically',
    ])

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 12 — SONARQUBE
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 12: SonarQube — Code Quality & Security')

    h2(doc, '12.1 Set Kernel Parameters (if not done in Part 5)')
    code(doc,
        'sudo sysctl -w vm.max_map_count=524288\n'
        'echo "vm.max_map_count=524288" | sudo tee -a /etc/sysctl.conf',
        'EC2 SSH'
    )

    h2(doc, '12.2 Run SonarQube via Docker')
    body(doc, 'The simplest installation — Docker must already be installed and the ubuntu user in docker group:')
    code(doc,
        'docker run -d \\\n'
        '  --name sonarqube \\\n'
        '  --restart unless-stopped \\\n'
        '  -p 9000:9000 \\\n'
        '  -e SONAR_ES_BOOTSTRAP_CHECKS_DISABLE=true \\\n'
        '  sonarqube:community\n\n'
        '# Watch startup logs (takes 2-3 minutes)\n'
        'docker logs -f sonarqube\n'
        '# Wait for: "SonarQube is operational"  then press Ctrl+C',
        'EC2 SSH'
    )
    note(doc, 'SonarQube needs at least 2 GB RAM free. Make sure no other heavy processes are competing.')

    h2(doc, '12.3 Fix Elasticsearch Disk Watermark (Run After Every Restart)')
    body(doc, 'If SonarQube analysis fails with "disk usage exceeded flood-stage watermark", run:')
    code(doc,
        '# Unlock read-only indices\n'
        'docker exec sonarqube curl -s -X PUT "http://localhost:9001/_all/_settings" \\\n'
        '  -H \'Content-Type: application/json\' \\\n'
        '  -d \'{"index.blocks.read_only_allow_delete": null}\'\n\n'
        '# Set lenient watermarks\n'
        'docker exec sonarqube curl -s -X PUT "http://localhost:9001/_cluster/settings" \\\n'
        '  -H \'Content-Type: application/json\' \\\n'
        '  -d \'{"persistent":{"cluster.routing.allocation.disk.watermark.low":"90%",\n'
        '       "cluster.routing.allocation.disk.watermark.high":"93%",\n'
        '       "cluster.routing.allocation.disk.watermark.flood_stage":"97%"}}\'\n\n'
        '# Then free disk space\n'
        'docker image prune -af\n'
        'docker builder prune -af',
        'EC2 SSH'
    )

    h2(doc, '12.4 Create SonarQube Project')
    steps(doc, [
        'Open http://<EC2-IP>:9000 → login admin/admin → change password',
        'Click Create a local project',
        'Project key: wahaj-nextjs-app | Display name: Wahaj Next.js App',
        'Click Set Up → Locally',
        'Token name: jenkins-token → Generate → COPY THE TOKEN (save it!)',
        'Select analysis method: Other (sonar-scanner) → OS: Linux',
    ])
    warn(doc, 'Copy the generated token immediately. You cannot view it again after closing the page.')

    h2(doc, '12.5 Configure Quality Gate (Block Insecure Code)')
    steps(doc, [
        'SonarQube top menu → Quality Gates',
        'Click "Sonar way" (default gate)',
        'Confirm condition: Security Rating is worse than A → this blocks deployment',
        'Go to your project → Project Settings → Quality Gate → assign "Sonar way"',
        'Click Set as Default if not already set',
    ])

    h2(doc, '12.6 sonar-project.properties')
    body(doc, 'Create this file in the root of your repository:')
    code(doc,
        'sonar.projectKey=wahaj-nextjs-app\n'
        'sonar.projectName=Wahaj Next.js App - FA22-BSE-100\n'
        'sonar.projectVersion=1.0\n'
        'sonar.sources=app\n'
        'sonar.exclusions=**/node_modules/**,**/.next/**,**/public/**\n'
        'sonar.sourceEncoding=UTF-8\n'
        'sonar.qualitygate.wait=true',
        'sonar-project.properties (repo root)'
    )

    h2(doc, '12.7 SonarLint — IDE Integration (VS Code)')
    steps(doc, [
        'VS Code → Extensions → search SonarLint → Install',
        'Command Palette (Ctrl+Shift+P) → SonarLint: Connect to SonarQube',
        'Server URL: http://<EC2-IP>:9000',
        'Generate user token: SonarQube → My Account → Security → Generate Token',
        'Paste token into VS Code SonarLint connection',
        'Select project: wahaj-nextjs-app',
        'SonarLint now shows issues inline as you type',
    ])

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 13 — JENKINS PIPELINE
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 13: Jenkins Pipeline — Full CI/CD Setup')

    h2(doc, '13.1 Add SonarQube Token to Jenkins Credentials')
    steps(doc, [
        'Jenkins → Manage Jenkins → Credentials → System → Global credentials → Add Credentials',
        'Kind: Secret text',
        'Secret: <paste your sonarqube token>',
        'ID: sonar-token',
        'Description: SonarQube Project Token',
        'Click Create',
    ])

    h2(doc, '13.2 Configure SonarQube Server in Jenkins')
    steps(doc, [
        'Manage Jenkins → System → scroll to SonarQube servers',
        'Check: Enable injection of SonarQube server configuration as environment variables',
        'Click Add SonarQube',
        'Name: SonarQube',
        'Server URL: http://localhost:9000',
        'Server authentication token: select sonar-token',
        'Click Save',
    ])

    h2(doc, '13.3 Add Jenkins Webhook in SonarQube')
    steps(doc, [
        'SonarQube → Administration → Configuration → Webhooks → Create',
        'Name: Jenkins',
        'URL: http://<EC2-IP>:8080/sonarqube-webhook/',
        'Leave Secret blank → Save',
    ])
    note(doc, 'Use the EC2 public IP, NOT localhost, in the SonarQube webhook URL. SonarQube blocks loopback addresses in webhooks.')

    h2(doc, '13.4 Create Jenkins Pipeline Job')
    steps(doc, [
        'Jenkins dashboard → New Item',
        'Name: wahaj-pipeline → select Pipeline → OK',
        'Scroll to Pipeline section',
        'Definition: Pipeline script from SCM',
        'SCM: Git',
        'Repository URL: https://github.com/ESE-Wahaj/wahajdockerjenkins',
        'Branch Specifier: */main   (NOT */master)',
        'Script Path: Jenkinsfile',
        'Click Save',
    ])

    h2(doc, '13.5 Jenkinsfile — Complete Pipeline')
    code(doc,
        'node {\n'
        '    def imageName    = \'wahaj-app\'\n'
        '    def containerName = \'wahaj-app\'\n'
        '    def appPort      = \'3001\'\n'
        '    def k8sNamespace = \'wahaj-app\'\n'
        '    def sonarProject = \'wahaj-nextjs-app\'\n\n'
        '    stage(\'Clean Workspace\') {\n'
        '        deleteDir()\n'
        '    }\n\n'
        '    stage(\'Clone Repo\') {\n'
        '        git(branch: \'main\',\n'
        '            url: \'https://github.com/ESE-Wahaj/wahajdockerjenkins\')\n'
        '    }\n\n'
        '    stage(\'SonarQube Analysis\') {\n'
        '        def scannerHome = tool \'SonarQube\'\n'
        '        withSonarQubeEnv(\'SonarQube\') {\n'
        '            sh """\n'
        '                ${scannerHome}/bin/sonar-scanner \\\n'
        '                  -Dsonar.projectKey=${sonarProject} \\\n'
        '                  -Dsonar.sources=app \\\n'
        '                  -Dsonar.token=${SONAR_AUTH_TOKEN}\n'
        '            """\n'
        '        }\n'
        '    }\n\n'
        '    stage(\'Quality Gate\') {\n'
        '        timeout(time: 5, unit: \'MINUTES\') {\n'
        '            def qg = waitForQualityGate()\n'
        '            if (qg.status != \'OK\') {\n'
        '                error "BLOCKED: Quality Gate: ${qg.status}"\n'
        '            }\n'
        '        }\n'
        '    }\n\n'
        '    stage(\'Build Image\') {\n'
        '        sh "docker build -t ${imageName}:${BUILD_NUMBER} ."\n'
        '        sh "docker tag ${imageName}:${BUILD_NUMBER} ${imageName}:latest"\n'
        '    }\n\n'
        '    stage(\'Deploy Docker Container\') {\n'
        '        sh """\n'
        '            docker stop ${containerName} || true\n'
        '            docker rm   ${containerName} || true\n'
        '            docker run -d --name ${containerName} \\\n'
        '              -p ${appPort}:3000 --restart unless-stopped \\\n'
        '              ${imageName}:${BUILD_NUMBER}\n'
        '        """\n'
        '    }\n\n'
        '    stage(\'Load Image into Minikube\') {\n'
        '        sh "minikube image load ${imageName}:latest"\n'
        '    }\n\n'
        '    stage(\'Deploy to Kubernetes\') {\n'
        '        sh "kubectl apply -f kubernetes/namespace.yaml"\n'
        '        sh "kubectl apply -f kubernetes/deployment.yaml"\n'
        '        sh "kubectl apply -f kubernetes/service.yaml"\n'
        '        sh "kubectl rollout restart deployment/${imageName} -n ${k8sNamespace}"\n'
        '        sh "kubectl rollout status  deployment/${imageName} -n ${k8sNamespace} --timeout=60s"\n'
        '    }\n\n'
        '    stage(\'Verify Deployment\') {\n'
        '        sh "kubectl get pods -n ${k8sNamespace}"\n'
        '        sh "kubectl get svc  -n ${k8sNamespace}"\n'
        '    }\n'
        '}',
        'Jenkinsfile'
    )

    h2(doc, '13.6 Run the Pipeline')
    steps(doc, [
        'Click Build Now on the left sidebar',
        'Click the build number → Console Output to watch live',
        'All 8 stages should go green — total time approx 3-5 minutes',
        'On success: app runs at http://<EC2-IP>:3001 (Docker) and http://<EC2-IP>:30080 (Kubernetes)',
    ])

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 14 — END TO END FLOW
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 14: Complete End-to-End Flow')
    body(doc, 'What happens from a git push to a running Kubernetes pod:')

    table2(doc,
        ['Step', 'Where', 'What Happens'],
        [
            ('1',  'VS Code',          'Developer writes code; SonarLint flags issues inline'),
            ('2',  'VS Code',          'git add . && git commit -m "..." && git push origin main'),
            ('3',  'GitHub',           'Code pushed to main branch; webhook fires to Jenkins'),
            ('4',  'Jenkins',          'Stage: Clean Workspace — old files wiped'),
            ('5',  'Jenkins',          'Stage: Clone Repo — latest code pulled from GitHub'),
            ('6',  'Jenkins→SonarQube','Stage: SonarQube Analysis — scanner analyses app/ folder'),
            ('7',  'SonarQube',        'Compute Engine processes report; evaluates Quality Gate'),
            ('8',  'Jenkins',          'Stage: Quality Gate — BLOCKS pipeline if Security Rating < A'),
            ('9',  'Jenkins',          'Stage: Build Image — docker build creates wahaj-app:N'),
            ('10', 'Jenkins',          'Stage: Deploy Docker — old container stopped; new one started on :3001'),
            ('11', 'Jenkins→Minikube', 'Stage: Load Image — minikube image load wahaj-app:latest'),
            ('12', 'Jenkins→K8s',      'Stage: Deploy K8s — kubectl apply manifests; rolling restart'),
            ('13', 'Jenkins',          'Stage: Verify — pod Running; NodePort service on :30080 confirmed'),
            ('14', 'Prometheus',       'Scrapes metrics every 15s from node_exporter and jenkins'),
            ('15', 'Grafana',          'Dashboards show live CPU, RAM, disk, Jenkins build stats'),
        ]
    )

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 15 — SERVICE URLS
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 15: Service URLs Reference')
    note(doc, 'Replace <EC2-IP> with your current EC2 public IP. It changes on every instance restart.')

    table2(doc,
        ['Service', 'URL', 'Default Login'],
        [
            ('Jenkins',             'http://<EC2-IP>:8080',          'admin / <set during setup>'),
            ('Web App (Docker)',     'http://<EC2-IP>:3001',          'N/A'),
            ('Web App (Kubernetes)','http://<EC2-IP>:30080',         'N/A'),
            ('SonarQube',           'http://<EC2-IP>:9000',          'admin / admin (change it!)'),
            ('Prometheus',          'http://<EC2-IP>:9090',          'N/A (no login)'),
            ('Grafana',             'http://<EC2-IP>:3030',          'admin / admin (change it!)'),
            ('Node Exporter',       'http://<EC2-IP>:9100/metrics',  'N/A (raw metrics)'),
        ]
    )

    pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 16 — TROUBLESHOOTING
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, 'Part 16: Troubleshooting')

    h2(doc, '16.1 Docker: permission denied')
    code(doc,
        'sudo usermod -aG docker ubuntu\n'
        'sudo usermod -aG docker jenkins\n'
        'newgrp docker\n'
        'sudo systemctl restart jenkins',
        'EC2 SSH'
    )

    h2(doc, '16.2 No space left on device')
    code(doc,
        'df -h                          # Check usage\n'
        'docker system prune -af        # Remove unused images/containers\n'
        'docker builder prune -af       # Clear build cache\n'
        'sudo rm -rf /var/lib/jenkins/workspace/*  # Clear Jenkins workspaces\n'
        'sudo apt-get clean             # Clear apt cache\n'
        'sudo journalctl --vacuum-size=100M',
        'EC2 SSH'
    )

    h2(doc, '16.3 Minikube cluster does not exist')
    code(doc,
        '# Start Minikube as jenkins user\n'
        'sudo su - jenkins -s /bin/bash -c \\\n'
        '  "minikube start --driver=docker --memory=2048 --cpus=2"\n\n'
        '# Check status\n'
        'sudo su - jenkins -s /bin/bash -c "minikube status"',
        'EC2 SSH'
    )

    h2(doc, '16.4 SonarQube CE Task FAILED')
    code(doc,
        '# Check logs\n'
        'docker logs sonarqube 2>&1 | grep -i "error\\|fail" | tail -20\n\n'
        '# Fix: unlock ES indices + set watermarks\n'
        'docker exec sonarqube curl -s -X PUT "http://localhost:9001/_all/_settings" \\\n'
        '  -H \'Content-Type: application/json\' \\\n'
        '  -d \'{"index.blocks.read_only_allow_delete": null}\'\n\n'
        'docker exec sonarqube curl -s -X PUT "http://localhost:9001/_cluster/settings" \\\n'
        '  -H \'Content-Type: application/json\' \\\n'
        '  -d \'{"persistent":{"cluster.routing.allocation.disk.watermark.flood_stage":"97%"}}\'\n\n'
        'docker restart sonarqube',
        'EC2 SSH'
    )

    h2(doc, '16.5 Jenkins branch error: master not found')
    body(doc, 'The repo uses "main" not "master". In the pipeline job:')
    steps(doc, [
        'wahaj-pipeline → Configure',
        'Branches to build → change */master to */main',
        'Save → Build Now',
    ])

    h2(doc, '16.6 sonar-scanner not found')
    body(doc, 'Jenkins must have SonarQube Scanner configured as a Tool:')
    steps(doc, [
        'Manage Jenkins → Tools → SonarQube Scanner → Add',
        'Name: SonarQube (must match exactly what is in Jenkinsfile)',
        'Check: Install automatically',
        'Save → run Build Now (Jenkins downloads it on first run)',
    ])

    h2(doc, '16.7 EC2 IP changed after restart')
    code(doc,
        '# Get new IP (run on EC2 or check AWS Console)\n'
        'curl http://checkip.amazonaws.com\n\n'
        '# Update inventory file\n'
        'nano ansible/inventory.ini   # change ansible_host=<NEW-IP>\n\n'
        '# Update SonarQube webhook in SonarQube UI\n'
        '# Update SonarLint connection in VS Code',
        'EC2 SSH'
    )

    h2(doc, '16.8 Corrupted Docker Layer Cache')
    code(doc,
        'docker builder prune -af\n'
        'docker system prune -f\n'
        '# Then retry the Jenkins build',
        'EC2 SSH'
    )

    # ── FOOTER ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('Muhammad Wahaj Naveed  |  FA22-BSE-100  |  Complete DevOps Setup Guide  |  May 2026')
    fr.font.size = Pt(9); fr.font.italic = True; fr.font.color.rgb = GRAY

    doc.save(OUTPUT)
    print(f'Saved: {OUTPUT}')

if __name__ == '__main__':
    build()
