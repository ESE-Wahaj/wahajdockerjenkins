"""
Run: pip install python-docx
Then: python docs/generate_docs.py
Outputs: docs/DevOps_FA22BSE100_Wahaj_Documentation.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "DevOps_FA22BSE100_Wahaj_Documentation.docx")

ORANGE = RGBColor(0xE8, 0x50, 0x00)
BLUE   = RGBColor(0x1F, 0x49, 0x7D)
GRAY   = RGBColor(0x36, 0x36, 0x36)
CODE_BG = RGBColor(0x1E, 0x1E, 0x1E)
CODE_FG = RGBColor(0x00, 0xFF, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.color.rgb = ORANGE if level == 1 else BLUE
    run.font.bold = True
    h.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    h.paragraph_format.space_after = Pt(4)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY
    return p


def add_code(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xD4, 0x00, 0x00)
    # light gray shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    return p


def add_info_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, 'FFF3CD')
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x85, 0x63, 0x04)
    doc.add_paragraph()


def add_url_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'URL'
    for r in hdr:
        for p in r.paragraphs:
            for run in p.runs:
                run.font.bold = True
    for label, url in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = url
    doc.add_paragraph()


def add_step_table(doc, steps):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Step'
    hdr[1].text = 'Action'
    for r in hdr:
        for p in r.paragraphs:
            for run in p.runs:
                run.font.bold = True
    for i, step in enumerate(steps, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = step
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
def build_document():
    doc = Document()

    # ── Margins ──────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover Page ───────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DevOps Full Pipeline Documentation")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = ORANGE

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run(
        "GitHub → Jenkins → SonarQube → Docker → Kubernetes\n"
        "Infrastructure as Code (Terraform) | Configuration Management (Ansible)\n"
        "Monitoring: Prometheus + Grafana"
    )
    run2.font.size = Pt(13)
    run2.font.color.rgb = BLUE

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = meta.add_run(
        "Muhammad Wahaj Naveed\nFA22-BSE-100\n"
        "EC2 IP: 54.234.94.152\nDate: May 2026"
    )
    run3.font.size = Pt(12)
    run3.font.color.rgb = GRAY

    page_break(doc)

    # ── Architecture Overview ─────────────────────────────────────────────────
    add_heading(doc, "1. Architecture Overview")
    add_body(doc,
        "The complete DevOps pipeline is hosted on a single AWS EC2 instance (t2.medium, Ubuntu 24.04). "
        "All components are provisioned via Terraform, configured via Ansible, and run as systemd services "
        "or Kubernetes pods inside Minikube."
    )
    add_code(doc,
        "GitHub Repo\n"
        "    │\n"
        "    ▼  (webhook / manual trigger)\n"
        "Jenkins :8080\n"
        "    │\n"
        "    ├─► SonarQube Analysis :9000\n"
        "    │       └─► Quality Gate (BLOCK if FAIL)\n"
        "    │\n"
        "    ├─► Docker Build → Image tagged :BUILD_NUMBER\n"
        "    │       ├─► Docker Container :3001\n"
        "    │       └─► Minikube → K8s Pod :30080 (namespace: wahaj-app)\n"
        "    │\n"
        "    └─► Prometheus :9090 ◄── Node Exporter :9100\n"
        "              └─► Grafana :3030"
    )

    add_heading(doc, "Service Port Map", level=2)
    add_url_table(doc, [
        ("Jenkins",              "http://54.234.94.152:8080"),
        ("Web App (Docker)",     "http://54.234.94.152:3001"),
        ("Web App (Kubernetes)", "http://54.234.94.152:30080"),
        ("SonarQube",            "http://54.234.94.152:9000"),
        ("Prometheus",           "http://54.234.94.152:9090"),
        ("Grafana",              "http://54.234.94.152:3030"),
        ("Node Exporter",        "http://54.234.94.152:9100/metrics"),
    ])

    page_break(doc)

    # ── Part 1: GitHub ────────────────────────────────────────────────────────
    add_heading(doc, "2. GitHub — Source Control Setup")
    add_body(doc,
        "The application source code is hosted on GitHub at "
        "https://github.com/ESE-Wahaj/wahajdockerjenkins. "
        "The repository contains the Next.js application, Dockerfile, Jenkinsfile, "
        "Terraform configurations, Ansible playbooks, and Kubernetes manifests."
    )

    add_heading(doc, "Repository Structure", level=2)
    add_code(doc,
        "wahajdockerjenkins/\n"
        "├── app/                        # Next.js app source\n"
        "├── terraform/                  # AWS infrastructure as code\n"
        "│   ├── main.tf\n"
        "│   ├── variables.tf\n"
        "│   └── outputs.tf\n"
        "├── ansible/                    # Configuration management\n"
        "│   ├── inventory.ini\n"
        "│   ├── playbook.yml\n"
        "│   └── sonarqube-playbook.yml\n"
        "├── kubernetes/                 # K8s manifests\n"
        "│   ├── namespace.yaml\n"
        "│   ├── deployment.yaml\n"
        "│   └── service.yaml\n"
        "├── docs/                       # This documentation\n"
        "├── Dockerfile                  # Multi-stage Docker build\n"
        "├── Jenkinsfile                 # CI/CD pipeline definition\n"
        "└── sonar-project.properties    # SonarQube project config"
    )

    add_heading(doc, "GitHub to Jenkins Webhook", level=2)
    add_body(doc,
        "Jenkins polls or receives webhook events from GitHub to trigger the pipeline "
        "automatically on every push to the main branch."
    )
    add_step_table(doc, [
        "In Jenkins: Manage Jenkins → Configure System → GitHub → Add GitHub Server",
        "Add GitHub personal access token as a Secret Text credential",
        "In your GitHub repo: Settings → Webhooks → Add webhook",
        "Payload URL: http://54.234.94.152:8080/github-webhook/",
        "Content type: application/json | Trigger: Just the push event",
        "Save — now every git push triggers Jenkins automatically",
    ])

    page_break(doc)

    # ── Part 2: Terraform ─────────────────────────────────────────────────────
    add_heading(doc, "3. Infrastructure as Code — Terraform")
    add_body(doc,
        "Terraform provisions the EC2 instance and security group on AWS. "
        "All infrastructure state is declared in HCL (HashiCorp Configuration Language) "
        "files under the terraform/ directory."
    )

    add_heading(doc, "Resources Created", level=2)
    add_body(doc, "• EC2 Instance: t2.medium, Ubuntu 24.04 LTS, 20 GB gp3 EBS")
    add_body(doc, "• Security Group: wahaj-devops-sg with the following ports open:")
    add_code(doc,
        "Port 22    — SSH\n"
        "Port 80    — HTTP\n"
        "Port 3000  — Next.js App\n"
        "Port 3001  — Next.js Alt\n"
        "Port 3030  — Grafana\n"
        "Port 8080  — Jenkins\n"
        "Port 9000  — SonarQube\n"
        "Port 9090  — Prometheus\n"
        "Port 9100  — Node Exporter\n"
        "Port 30000-32767 — Kubernetes NodePorts"
    )

    add_heading(doc, "How to Run", level=2)
    add_code(doc,
        "# Install Terraform: https://developer.hashicorp.com/terraform/install\n\n"
        "cd terraform\n"
        "terraform init          # Download AWS provider\n"
        "terraform plan          # Preview resources\n"
        "terraform apply         # Create EC2 + security group\n"
        "terraform output        # Print all URLs\n"
        "terraform destroy       # Tear down when done"
    )
    add_info_box(doc,
        "NOTE: Set your AWS credentials before running Terraform:\n"
        "aws configure   (enter Access Key ID, Secret Key, region: us-east-1)"
    )

    page_break(doc)

    # ── Part 3: Ansible ───────────────────────────────────────────────────────
    add_heading(doc, "4. Configuration Management — Ansible")
    add_body(doc,
        "Ansible automates the installation and configuration of all services on the EC2 instance. "
        "Two playbooks are provided: one for the base stack (Docker, Minikube, Prometheus, Grafana) "
        "and one specifically for SonarQube."
    )

    add_heading(doc, "Inventory File (ansible/inventory.ini)", level=2)
    add_code(doc,
        "[ec2]\n"
        "wahaj-server ansible_host=54.234.94.152 \\\n"
        "  ansible_user=ubuntu \\\n"
        "  ansible_ssh_private_key_file=~/.ssh/jenkey.pem \\\n"
        "  ansible_ssh_common_args='-o StrictHostKeyChecking=no'"
    )
    add_info_box(doc,
        "IMPORTANT: The EC2 public IP changes on every instance restart. "
        "Always update ansible_host= in inventory.ini with the new IP before running."
    )

    add_heading(doc, "Run Base Stack Playbook", level=2)
    add_code(doc,
        "# Test connectivity\n"
        "ansible -i ansible/inventory.ini ec2 -m ping\n\n"
        "# Install Docker, kubectl, Minikube, Prometheus, Node Exporter, Grafana\n"
        "ansible-playbook -i ansible/inventory.ini ansible/playbook.yml"
    )

    add_heading(doc, "Run SonarQube Playbook", level=2)
    add_code(doc,
        "ansible-playbook -i ansible/inventory.ini ansible/sonarqube-playbook.yml"
    )
    add_body(doc,
        "The SonarQube playbook installs: Java 17, PostgreSQL, SonarQube 10.4 Community Edition, "
        "and SonarScanner CLI. It also sets the required kernel parameters (vm.max_map_count=524288) "
        "and creates the sonarqube systemd service."
    )

    page_break(doc)

    # ── Part 4: Docker ────────────────────────────────────────────────────────
    add_heading(doc, "5. Containerization — Docker")
    add_body(doc,
        "The Next.js application is packaged as a Docker image using a multi-stage build. "
        "The final image is based on node:20-alpine and runs as a non-root user for security."
    )

    add_heading(doc, "Dockerfile Stages", level=2)
    add_code(doc,
        "Stage 1 — deps   (node:20-alpine)\n"
        "  COPY package.json && RUN npm install\n\n"
        "Stage 2 — builder (node:20-alpine)\n"
        "  COPY source + node_modules\n"
        "  RUN npm run build  →  .next/standalone/\n\n"
        "Stage 3 — runner  (node:20-alpine)   ← Final image\n"
        "  Non-root user: nextjs (uid 1001)\n"
        "  EXPOSE 3000\n"
        "  CMD [\"node\", \"server.js\"]"
    )

    add_heading(doc, "next.config.ts — Standalone Mode", level=2)
    add_body(doc,
        "The Next.js app must use output: 'standalone' so the Docker build produces "
        "a self-contained server.js bundle without needing node_modules at runtime."
    )
    add_code(doc,
        "// next.config.ts\n"
        "const nextConfig = { output: 'standalone' };\n"
        "export default nextConfig;"
    )

    add_heading(doc, "Build and Run Manually", level=2)
    add_code(doc,
        "# Build image\n"
        "docker build -t wahaj-app:latest .\n\n"
        "# Run container (maps port 3001 → internal 3000)\n"
        "docker run -d --name wahaj-app -p 3001:3000 \\\n"
        "  --restart unless-stopped wahaj-app:latest\n\n"
        "# View logs\n"
        "docker logs -f wahaj-app\n\n"
        "# App is live at http://54.234.94.152:3001"
    )

    page_break(doc)

    # ── Part 5: Kubernetes ────────────────────────────────────────────────────
    add_heading(doc, "6. Container Orchestration — Kubernetes (Minikube)")
    add_body(doc,
        "The application runs inside a Kubernetes pod managed by Minikube. "
        "Minikube uses the Docker driver so no separate VM is needed on the EC2 instance."
    )

    add_heading(doc, "Manifest Files", level=2)
    add_code(doc,
        "kubernetes/namespace.yaml   — Creates namespace 'wahaj-app'\n"
        "kubernetes/deployment.yaml  — Pod spec: 1 replica, resource limits, health probes\n"
        "kubernetes/service.yaml     — NodePort service: 30080 → pod:3000"
    )

    add_heading(doc, "Key Deployment Settings", level=2)
    add_code(doc,
        "Namespace:        wahaj-app\n"
        "Replicas:         1\n"
        "Container port:   3000\n"
        "NodePort:         30080  (external)\n"
        "imagePullPolicy:  Never  (uses Minikube local cache)\n"
        "Memory limit:     512Mi\n"
        "CPU limit:        500m\n"
        "Readiness probe:  GET / on port 3000, delay 10s\n"
        "Liveness probe:   GET / on port 3000, delay 15s"
    )

    add_heading(doc, "Manual Deployment Steps", level=2)
    add_step_table(doc, [
        "minikube start --driver=docker --memory=2048 --cpus=2",
        "docker build -t wahaj-app:latest .",
        "minikube image load wahaj-app:latest",
        "kubectl apply -f kubernetes/namespace.yaml",
        "kubectl apply -f kubernetes/deployment.yaml",
        "kubectl apply -f kubernetes/service.yaml",
        "kubectl get pods -n wahaj-app  (wait for Running status)",
        "Access app at http://54.234.94.152:30080",
    ])

    add_heading(doc, "Useful Commands", level=2)
    add_code(doc,
        "kubectl get pods -n wahaj-app\n"
        "kubectl logs -f deployment/wahaj-app -n wahaj-app\n"
        "kubectl describe pod -l app=wahaj-app -n wahaj-app\n"
        "kubectl rollout restart deployment/wahaj-app -n wahaj-app\n"
        "minikube dashboard  (opens browser-based K8s UI)"
    )

    page_break(doc)

    # ── Part 6: SonarQube ─────────────────────────────────────────────────────
    add_heading(doc, "7. Code Quality & Security — SonarQube + SonarLint")
    add_body(doc,
        "SonarQube performs static code analysis to detect bugs, code smells, security "
        "vulnerabilities, and duplicated code. A Quality Gate is configured to BLOCK "
        "deployment if any security issues or quality thresholds are violated."
    )

    add_heading(doc, "7.1 Install SonarQube on EC2", level=2)
    add_body(doc,
        "Run the Ansible SonarQube playbook (see Part 4). It installs Java 17, PostgreSQL, "
        "and SonarQube 10.4 Community Edition as a systemd service on port 9000."
    )
    add_code(doc,
        "ansible-playbook -i ansible/inventory.ini ansible/sonarqube-playbook.yml\n\n"
        "# SonarQube starts at: http://54.234.94.152:9000\n"
        "# Default credentials: admin / admin  (change on first login)"
    )
    add_info_box(doc,
        "SonarQube takes 2-3 minutes to start after the playbook completes. "
        "It uses PostgreSQL as its database and Elasticsearch (embedded) for search. "
        "The EC2 instance must have at least 2 GB RAM (t2.medium recommended)."
    )

    add_heading(doc, "7.2 Create SonarQube Project", level=2)
    add_step_table(doc, [
        "Open http://54.234.94.152:9000 and log in as admin",
        "Click 'Create Project' → 'Manually'",
        "Project key: wahaj-nextjs-app | Display name: Wahaj Next.js App",
        "Click 'Set Up' → choose 'Locally'",
        "Generate a project token — copy and save it (you'll need it for Jenkins)",
        "Select analysis method: Other (sonar-scanner)",
    ])

    add_heading(doc, "7.3 Configure Quality Gate (Block Insecure Code)", level=2)
    add_body(doc,
        "The Quality Gate defines pass/fail conditions for code analysis. "
        "The default 'Sonar way' gate already blocks on Security Rating worse than A. "
        "To customize:"
    )
    add_step_table(doc, [
        "In SonarQube: Quality Gates → Create (name: Wahaj Security Gate)",
        "Add condition: Security Rating is worse than A  → BLOCKS build",
        "Add condition: Reliability Rating is worse than A  → BLOCKS build",
        "Add condition: Maintainability Rating is worse than A  → BLOCKS build",
        "Add condition: Coverage < 0%  (or your threshold)  → BLOCKS build",
        "Click 'Set as Default' so your project uses this gate",
        "In your project: Project Settings → Quality Gate → select Wahaj Security Gate",
    ])

    add_heading(doc, "7.4 sonar-project.properties", level=2)
    add_body(doc, "This file in the project root tells the scanner what to analyze:")
    add_code(doc,
        "sonar.projectKey=wahaj-nextjs-app\n"
        "sonar.projectName=Wahaj Next.js App - FA22-BSE-100\n"
        "sonar.projectVersion=1.0\n"
        "sonar.sources=app\n"
        "sonar.exclusions=**/node_modules/**,**/.next/**,**/public/**\n"
        "sonar.sourceEncoding=UTF-8\n"
        "sonar.qualitygate.wait=true"
    )

    add_heading(doc, "7.5 SonarLint (IDE Integration)", level=2)
    add_body(doc,
        "SonarLint is a VS Code extension that runs the same rules locally as you type, "
        "catching issues before you even push to GitHub."
    )
    add_step_table(doc, [
        "In VS Code: Extensions → search 'SonarLint' → Install",
        "Open Command Palette (Ctrl+Shift+P) → SonarLint: Connect to SonarQube",
        "Server URL: http://54.234.94.152:9000",
        "Generate a user token in SonarQube: My Account → Security → Generate Token",
        "Paste token into SonarLint connection setup",
        "Select project: wahaj-nextjs-app",
        "SonarLint now highlights issues inline in the editor",
    ])

    page_break(doc)

    # ── Part 7: Jenkins Pipeline ──────────────────────────────────────────────
    add_heading(doc, "8. CI/CD Pipeline — Jenkins")
    add_body(doc,
        "Jenkins orchestrates the full pipeline. Every push to GitHub triggers "
        "the Jenkinsfile which runs all stages in sequence. If the SonarQube "
        "Quality Gate fails, the pipeline is aborted and Docker/Kubernetes deployment never runs."
    )

    add_heading(doc, "8.1 Configure SonarQube in Jenkins", level=2)
    add_step_table(doc, [
        "Jenkins → Manage Jenkins → Plugins → Available: install 'SonarQube Scanner'",
        "Jenkins → Manage Jenkins → Plugins → Available: install 'Sonar Quality Gates'",
        "Jenkins → Manage Jenkins → Configure System → SonarQube Servers → Add",
        "Name: SonarQube | URL: http://localhost:9000",
        "Add credential: Secret Text → paste SonarQube project token → ID: sonar-token",
        "Jenkins → Manage Jenkins → Global Tool Configuration → SonarQube Scanner → Add",
        "Name: SonarQube | Install automatically (latest version)",
        "Save all settings",
    ])

    add_heading(doc, "8.2 Pipeline Stages", level=2)
    add_code(doc,
        "Stage 1: Clean Workspace       — wipe old files\n"
        "Stage 2: Clone Repo            — git pull from GitHub main\n"
        "Stage 3: SonarQube Analysis    — run sonar-scanner against app/ directory\n"
        "Stage 4: Quality Gate          — wait for result; ABORT if FAIL\n"
        "Stage 5: Build Image           — docker build + tag :BUILD_NUMBER + :latest\n"
        "Stage 6: Deploy Docker         — stop old container, run new one on :3001\n"
        "Stage 7: Load into Minikube    — minikube image load wahaj-app:latest\n"
        "Stage 8: Deploy to Kubernetes  — kubectl apply manifests + rollout restart\n"
        "Stage 9: Verify Deployment     — kubectl get pods + svc"
    )

    add_heading(doc, "8.3 Jenkinsfile SonarQube Snippet", level=2)
    add_code(doc,
        "stage('SonarQube Analysis') {\n"
        "    withSonarQubeEnv('SonarQube') {\n"
        "        sh 'sonar-scanner -Dsonar.projectKey=wahaj-nextjs-app ...'\n"
        "    }\n"
        "}\n\n"
        "stage('Quality Gate') {\n"
        "    timeout(time: 5, unit: 'MINUTES') {\n"
        "        def qg = waitForQualityGate()\n"
        "        if (qg.status != 'OK') {\n"
        "            error 'BLOCKED: Quality Gate failed: ' + qg.status\n"
        "        }\n"
        "    }\n"
        "}"
    )
    add_info_box(doc,
        "The Quality Gate stage uses a Jenkins webhook from SonarQube to get the result. "
        "In SonarQube: Administration → Configuration → Webhooks → Create:\n"
        "Name: Jenkins | URL: http://localhost:8080/sonarqube-webhook/"
    )

    page_break(doc)

    # ── Part 8: Monitoring ────────────────────────────────────────────────────
    add_heading(doc, "9. Monitoring — Prometheus + Grafana")
    add_body(doc,
        "Prometheus scrapes metrics from the EC2 system (via Node Exporter) and from Jenkins. "
        "Grafana visualizes these metrics on pre-built dashboards."
    )

    add_heading(doc, "Prometheus Scrape Targets", level=2)
    add_code(doc,
        "# /etc/prometheus/prometheus.yml\n"
        "scrape_configs:\n"
        "  - job_name: prometheus      # self-monitoring\n"
        "    targets: [localhost:9090]\n\n"
        "  - job_name: node_exporter   # CPU, RAM, Disk, Network\n"
        "    targets: [localhost:9100]\n\n"
        "  - job_name: jenkins         # Jenkins build metrics\n"
        "    metrics_path: /prometheus\n"
        "    targets: [localhost:8080]"
    )

    add_heading(doc, "Grafana Dashboards", level=2)
    add_code(doc,
        "Dashboard ID 1860  — Node Exporter Full (server metrics)\n"
        "Dashboard ID 9964  — Jenkins: Performance and Health Overview\n\n"
        "Import: Grafana → + → Import → enter ID → Load → select Prometheus data source"
    )

    page_break(doc)

    # ── Part 9: Complete Flow ─────────────────────────────────────────────────
    add_heading(doc, "10. End-to-End Flow Summary")
    add_body(doc,
        "The following sequence describes the complete pipeline from a developer "
        "writing code to the app running in a Kubernetes pod:"
    )
    add_step_table(doc, [
        "Developer writes code; SonarLint flags issues inline in VS Code",
        "Developer fixes issues and runs git push origin main",
        "GitHub sends webhook to Jenkins at port 8080",
        "Jenkins: Stage 1+2 — Clean workspace, clone latest code",
        "Jenkins: Stage 3 — sonar-scanner analyses the app/ directory",
        "SonarQube processes analysis and evaluates Quality Gate rules",
        "Jenkins: Stage 4 — waitForQualityGate(); ABORT if Security/Quality rules fail",
        "Jenkins: Stage 5 — docker build creates a new image tagged :BUILD_NUMBER",
        "Jenkins: Stage 6 — old Docker container stopped; new one started on port 3001",
        "Jenkins: Stage 7 — image loaded into Minikube image cache",
        "Jenkins: Stage 8 — kubectl applies manifests; rolling restart of K8s pod",
        "Jenkins: Stage 9 — pod status verified; service NodePort confirmed at :30080",
        "Prometheus scrapes metrics; Grafana dashboards show live health data",
    ])

    page_break(doc)

    # ── Troubleshooting ───────────────────────────────────────────────────────
    add_heading(doc, "11. Troubleshooting")

    add_heading(doc, "SonarQube won't start", level=2)
    add_code(doc,
        "# Check logs\n"
        "sudo journalctl -u sonarqube -f\n"
        "sudo cat /opt/sonarqube/logs/sonar.log\n"
        "sudo cat /opt/sonarqube/logs/es.log\n\n"
        "# Most common fix — check vm.max_map_count\n"
        "sysctl vm.max_map_count  # must be >= 524288\n"
        "sudo sysctl -w vm.max_map_count=524288"
    )

    add_heading(doc, "Jenkins 'sonar-scanner not found'", level=2)
    add_code(doc,
        "# Verify scanner is installed\n"
        "sonar-scanner --version\n\n"
        "# If not found — install manually\n"
        "sudo ln -s /opt/sonar-scanner-*/bin/sonar-scanner /usr/local/bin/sonar-scanner"
    )

    add_heading(doc, "Minikube/kubectl errors", level=2)
    add_code(doc,
        "# Check minikube status\n"
        "minikube status\n\n"
        "# Restart minikube\n"
        "minikube stop && minikube start --driver=docker\n\n"
        "# Rebuild and reload image after every code change\n"
        "docker build -t wahaj-app:latest .\n"
        "minikube image load wahaj-app:latest\n"
        "kubectl rollout restart deployment/wahaj-app -n wahaj-app"
    )

    add_heading(doc, "EC2 IP changed after restart", level=2)
    add_code(doc,
        "# Get new IP from AWS Console or:\n"
        "curl http://checkip.amazonaws.com\n\n"
        "# Update inventory\n"
        "nano ansible/inventory.ini   # change ansible_host=<NEW_IP>\n\n"
        "# Update SonarLint connection in VS Code with new IP"
    )

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        "Muhammad Wahaj Naveed | FA22-BSE-100 | DevOps Final Exam | May 2026"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = GRAY
    footer_run.font.italic = True

    doc.save(OUTPUT_PATH)
    print(f"Document saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
